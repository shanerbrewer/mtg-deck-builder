from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("", 0)
title_run = title.add_run("BANANA-COIN: The Peel-to-Earn Revolution")
title_run.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("A Totally Serious Business Proposal for Making Obscene Amounts of Money from Bananas")
subtitle_run.italic = True
subtitle_run.font.size = Pt(13)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Executive Summary
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "The global banana market generates over $25 billion annually. Yet somehow, nobody has "
    "thought to disrupt it with blockchain technology, NFTs, and a subscription model. "
    "Until now. Introducing BANANA-COIN: the world's first Peel-to-Earn cryptocurrency "
    "ecosystem, powered by actual bananas, irrational optimism, and a surprisingly robust "
    "Discord server."
)

# The Problem
doc.add_heading("The Problem", level=1)
doc.add_paragraph(
    "Every day, millions of bananas go uneaten. They sit on counters, browning slowly, "
    "judging us. Meanwhile, crypto whales are desperately searching for their next "
    "investment. These two crises have never been solved simultaneously — until now."
)

# The Solution
doc.add_heading("The Solution: Peel-to-Earn (P2E)", level=1)
doc.add_paragraph(
    "Here's the plan, which is definitely not a pyramid scheme:"
)

steps = [
    ("Step 1 — Mint the Peel",
     "Each banana peel is photographed, assigned a unique cryptographic hash (the 'PeelHash'), "
     "and minted as an NFT on the BananChain. Rarer peels (spotted, curled tip, 'The Double Banana') "
     "command higher prices. A perfectly brown peel? That's a legendary drop."),
    ("Step 2 — The BananaMiner App",
     "Users download the BananaMiner app and scan their bananas daily. The app uses "
     "AI-powered ripeness detection to assign a Ripeness Score (RS). The riper the banana, "
     "the more BANANA-COIN you earn. This incentivizes users to let bananas ripen, "
     "reducing food waste and increasing engagement metrics simultaneously."),
    ("Step 3 — Staking & the Bunch Pool",
     "Holders can stake their BANANA-COIN in the 'Bunch Pool' to earn passive income. "
     "For every 12 coins staked (one bunch), users receive a weekly 'Smoothie Dividend' — "
     "paid in BANANA-COIN, naturally. APY is projected at 47%, a number we chose because "
     "it sounds believable but also exciting."),
    ("Step 4 — The Peel Marketplace",
     "A peer-to-peer marketplace where banana peel NFTs are traded. Premium tiers include: "
     "Standard Peel ($4.99), Artisanal Organic Peel ($24.99), and the ultra-rare "
     "Triple-Curved Heritage Peel (price negotiable, inquire within)."),
    ("Step 5 — Enterprise B2B: The Slip Licensing Program",
     "License the classic banana-slip gag to Hollywood studios, insurance companies (ironic "
     "marketing), and law firms specializing in personal injury. BANANA-COIN is accepted as payment. "
     "This is our 'real revenue' stream that makes the whole thing technically legal."),
]

for title_text, body_text in steps:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(body_text)

# Tokenomics
doc.add_heading("Tokenomics", level=1)
doc.add_paragraph("Total supply: 21,000,000 BANANA-COIN (because Bitcoin vibes)\n"
                  "Reserved for founders: 40% (we call it the 'Stem Allocation')\n"
                  "Community rewards: 30%\n"
                  "Marketing budget: 20% (mostly for monkey influencers)\n"
                  "Emergency peel reserve: 10%")

# Risk Factors
doc.add_heading("Risk Factors", level=1)
risks = [
    "Seasonal banana shortages may cause volatility.",
    "Monkeys. We have not fully modeled the monkey threat.",
    "Regulatory agencies may not recognize banana peels as securities. (We believe they are wrong.)",
    "If the price of BANANA-COIN falls below the price of an actual banana, morale will suffer.",
    "Our lead developer is allergic to bananas and has never touched one.",
]
for risk in risks:
    doc.add_paragraph(risk, style="List Bullet")

# Roadmap
doc.add_heading("Roadmap", level=1)
roadmap = [
    ("Q1 2026", "Launch BananaMiner app (iOS only — Android users can wait, they're used to it)"),
    ("Q2 2026", "First Peel NFT drop. Introduce the 'Banana DAO' governance token."),
    ("Q3 2026", "Partner with at least one grocery chain, or just hang a poster there without permission."),
    ("Q4 2026", "Launch BANANA-COIN on a major exchange, or a minor one, or just our own website."),
    ("2027",    "World domination. Or a pivot to mango. Jury's still out."),
]
for quarter, milestone in roadmap:
    p = doc.add_paragraph()
    p.add_run(f"{quarter}: ").bold = True
    p.add_run(milestone)

# Closing
doc.add_heading("Why This Will Work", level=1)
doc.add_paragraph(
    "Bananas are the world's most popular fruit. Crypto is the world's most popular speculation vehicle. "
    "Combining them is simply good business sense. We are not financial advisors. This document is for "
    "entertainment purposes only. Please do not actually invest in BANANA-COIN. But if you do, "
    "our wallet address is available on request."
)

p = doc.add_paragraph()
p.add_run("\n\nTo the moon. 🍌").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("banana_monetization_proposal.docx")
print("Document created: banana_monetization_proposal.docx")
